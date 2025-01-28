import os
from PIL import Image 
from docx import Document
from docx.shared import Inches, Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert


def page_size(doc):
    section = doc.sections[0]
    # Define A4 size
    width = Inches(8.27)
    height = Inches(11.69)

    section.page_width = width
    section.page_height = height

doc = Document()
page_size(doc)

section = doc.sections[0]
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.5)
# section.top_margin = Inches(0.19)
section.bottom_margin = Inches(0.38)
section.header_distance = Inches(0.22)


section.footer_distance=Inches(0.1)

head_table = section.header.add_table(rows=1, cols=2, width=Inches(7.27))
head_table.autofit = False  # Necessary if you want to set width manually
# add_bottom_margin_to_table(head_table, 0.1)  # 0.1 inches bottom margin




#clear by deafult paragraph 
cell1 = head_table.cell(0,0)
cell1._element.clear_content() 



# Set cell widths
head_table.cell(0, 0).width = Inches(2.9)
head_table.cell(0, 1).width = Inches(4.37)

# Set row height
for row in head_table.rows:
    row.height = Pt(54)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

# Add paragraph and run to the cell
para1 = head_table.cell(0,0).add_paragraph()
run = para1.add_run()
para1.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Add picture to the cell
para1.add_run().add_picture("D:\Btech\s3\Python Programming\docx_learning_miniProject\marwadi.png", width=Inches(2.9), height=Inches(0.75))



# Add text to the second column and align it to the right
cell_text = head_table.cell(0, 1)
cell_text._element.clear_content()
paragraph_text = cell_text.add_paragraph()
paragraph_text.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Right alignment

# Add the three lines of text
run = paragraph_text.add_run("FACULTY OF ENGINEERING AND TECHNOLOGY")
run.add_break()
run.font.name = "Calibri (Body)"
run.bold = True
run.font.size = Pt(14)

run = paragraph_text.add_run("Department of Computer Engineering")
run.add_break()
# run.bold = True
run.font.name = "Calibri (Body)"
run.font.size = Pt(12)

run = paragraph_text.add_run("01CE0412 –Advance Web Technology – Lab Manual")
run.font.name = "Calibri (Body)"
run.font.size = Pt(12)


# Apply a consistent bottom border to both cells
for cell in head_table.row_cells(0):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '13')  # Thickness of the line
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


#now FOOTER sectoin begins
# Now FOOTER section begins
# Create and format the footer table
footer_table = section.footer.add_table(rows=1, cols=2, width=Inches(7.27))
footer_table.autofit = False

# Get user input for starting page number
start_page = int(input("Enter the starting page number: "))

# Set column widths
footer_table.cell(0, 0).width = Inches(6.5)  # Adjust width as needed
footer_table.cell(0, 1).width = Inches(0.77)  # Adjust width as needed

# Set row height
for row in footer_table.rows:
    row.height = Pt(0.31 * 72)  # Convert 0.31 inches to points (72 points per inch)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

# Add content to the first cell (left alignment)
cell_text = footer_table.cell(0, 0)
cell_text._element.clear_content()
ptext = cell_text.add_paragraph()
ptext.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


# ---------------------------------------------------------------enrollment--------------------------------------
# ---------------------------------------------------------------enrollment--------------------------------------
# ---------------------------------------------------------------enrollment--------------------------------------
# ---------------------------------------------------------------enrollment--------------------------------------
# ---------------------------------------------------------------enrollment--------------------------------------

ptext.add_run("92410103123")




# Add page number field to the second cell (right alignment)
cell_text = footer_table.cell(0, 1)
cell_text._element.clear_content()
ptext = cell_text.add_paragraph()
ptext.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Create the page number field
fldSimple = OxmlElement('w:fldSimple')
fldSimple.set(qn('w:instr'), 'PAGE')  # Insert page number field

# Append the field to the paragraph
ptext.add_run("| ")
ptext._p.append(fldSimple)

# Adjust the section properties to start from the specified page number
# Access section properties directly from the section object
sectPr = section._sectPr
pgNumType = OxmlElement('w:pgNumType')
pgNumType.set(qn('w:start'), str(start_page))
sectPr.append(pgNumType)


# Apply a consistent bottom border to both cells
for cell in footer_table.row_cells(0):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    bottom = OxmlElement('w:top')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '13')  # Thickness of the line
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


#now begin to copy and paste one files code to here.
heads = []
def_head = input("enter practicals defination : ")
ask = input("a,b,c,d,e,f any sub defination (if not then write NO): ")
c = 0
if ask != "NO".lower():
    heads.append(def_head)
    heads.append(ask)
    p_head_def = doc.add_paragraph()
    for head in heads:
        run = p_head_def.add_run(head)
        if c == 0:
            run.add_break()
            c += 1
        run.bold = True
        run.font.name = "Calibri (Body)"
        run.font.size = Pt(12)

    run.add_break()
else:
    heads.append(def_head)
    p_head_def = doc.add_paragraph()
    run = p_head_def.add_run(def_head) 
    run.bold = True
    run.font.name = "Calibri (Body)"
    run.font.size = Pt(12)

# File paths
html_file_path = input("enter file path : ").strip('"')
c=0
for i in range(len(html_file_path)-1,-1,-1):
    if html_file_path[i] == "\\":
        break
    else:
        c+=1
main_path = html_file_path[:-c]
  
    
# file_name = input("enter file name : ")
if not main_path.endswith("\\"):
    main_path += "\\"

image_list=[]
for i in os.listdir(main_path):
    if i.lower().endswith(".png") or i.lower().endswith(".jpg") or i.lower().endswith(".jpeg"):
        image_list.append(i)

with open(html_file_path,"r", encoding="utf-8") as file:
    data = file.readlines()

writer = doc.add_paragraph()
for line in data:
    run = writer.add_run(line.rstrip())
    run.font.name="Courier New"
    run.font.size = Pt(11)
    run.add_break()


for i in image_list:
    img_path = os.path.join(main_path, i.strip('"'))
    
    # Open the image to check its size
    with Image.open(img_path) as img:
        width_in_inches = img.width / img.info['dpi'][0] if 'dpi' in img.info else img.width / 96
        height_in_inches = img.height / img.info['dpi'][1] if 'dpi' in img.info else img.height / 96

    # Decide whether to resize or not
    if width_in_inches > 7 or height_in_inches > 5:
        doc.add_picture(img_path, width=Inches(7), height=Inches(3.5))
    else:
        doc.add_picture(img_path)

    # Add the image name
    img_name = doc.add_paragraph()
    img_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = img_name.add_run(i.strip('"'))
    run.font.name = "Calibri (Body)"
    run.font.size = Pt(12)
    run.add_break()
    
# Save the document 


# Define the file paths
docx_file_path = html_file_path.replace(".html", ".docx")
pdf_path = docx_file_path.replace(".docx", ".pdf")

try:
    # Save the document as a Word file
    if os.path.exists(docx_file_path):
        os.remove(docx_file_path)
    doc.save(docx_file_path)

    # Convert the Word document to a PDF file
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
        print("PDF file already exists, and it has been removed.")
    convert(docx_file_path, pdf_path)

    print("Conversion completed successfully.")
except Exception as e:
    print(f"An error occurred: {e}")